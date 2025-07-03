import numpy as np
import matplotlib.pyplot as plt
import scipy.fftpack
import soundfile as sf

import sounddevice as sd

from docx import Document
from docx.shared import Inches
from io import BytesIO

def plotAudio(Signal,Fs,TimeMargin=[0,0.02],fsize=2**8, axs=None):
    # plt.figure() 
    # plt.subplot(2,1,1) 
    # plt.plot(np.arange(0,Signal.shape[0])/Fs,Signal) 
    # plt.xlim(TimeMargin)
    # plt.subplot(2,1,2)
    # yf = scipy.fftpack.fft(Signal,fsize) 
    # plt.plot(np.arange(0,Fs,1.0*Fs/(yf.size)),np.abs(yf)) 
    # plt.show()
    
    # plt.figure() 
    # plt.subplot(2,1,1) 
    # plt.plot(np.arange(0,Signal.shape[0])/Fs,Signal) 
    # plt.subplot(2,1,2) 
    # yf = scipy.fftpack.fft(Signal,fsize) 
    # plt.plot(np.arange(0,Fs/2,Fs/fsize),20*np.log10( np.abs(yf[:fsize//2]))) 
    # plt.show()
    
    if axs is None:
        fig, axs = plt.subplots(2, 1, figsize=(10, 7))
    
    time = np.arange(0, len(Signal)) / Fs  
    
    yf = scipy.fftpack.fft(Signal, fsize)
    magnitude = np.abs(yf[:fsize//2])  
    magnitude_db = 20 * np.log10(magnitude + 1e-10)
    freqs = np.linspace(0, Fs/2, fsize//2)  
    max_index = np.argmax(magnitude)  
    max_freq = freqs[max_index]  
    max_amplitude = magnitude_db[max_index]  

    axs[0].plot(time, Signal, label='Sygnał audio', color='blue')
    axs[0].set_xlabel('Czas (s)')
    axs[0].set_ylabel('Amplituda (db)')
    axs[0].set_xlim(TimeMargin)
    axs[0].legend()
    
    axs[1].plot(freqs, magnitude_db, label='Widmo (dB)', color='red')
    axs[1].set_xlabel('Częstotliwość (Hz)')
    axs[1].set_ylabel('Amplituda (dB)')
    axs[1].legend()
    
    if axs is None:
        plt.show()
    return max_freq, max_amplitude

# Zad 1

data, fs = sf.read('sound1.wav', dtype='float32')

sf.write('sound_L.wav', data[:, 0], fs)  
sf.write('sound_R.wav', data[:, 1], fs)
sf.write('sound_mix.wav', (data[:, 0] + data[:, 1]) / 2, fs) 

# Zad 2

Signal, Fs = sf.read('sin_440Hz.wav', dtype=np.int32)
plotAudio(Signal,Fs,TimeMargin=[0, 0.02])

# Zad 3

document = Document()
document.add_heading('Lab1', 0)

files = ['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav']
fsize_values = [2**8, 2**12, 2**16]


for file in files:
    document.add_heading('Plik - {}'.format(file),2)
    data, fs = sf.read(file, dtype='float32')
    for fsize in fsize_values:
        document.add_heading(f'FFT Rozmiar {fsize}', 3)
        fig, axs = plt.subplots(2, 1, figsize=(10, 7))
        max_freq, max_amplitude = plotAudio(data, fs, TimeMargin=[0, 0.02], fsize=fsize, axs=axs)
        fig.suptitle(f'FFT Rozmiar {fsize}')
        fig.tight_layout(pad=1.5)
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        document.add_paragraph(f'Maksymalna częstotliwość: {max_freq:.2f} Hz')
        document.add_paragraph(f'Maksymalna amplituda: {max_amplitude:.2f} dB')
        plt.close(fig)

report_name = 'report.docx'
document.save(report_name)